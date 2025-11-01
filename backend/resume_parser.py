import re
import pdfplumber
import docx
import spacy
import os
import logging
from nltk.corpus import stopwords
import nltk

# Download resources silently
nltk.download('stopwords', quiet=True)

# Load NLP model
try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Common skills (you can expand or load dynamically from job DB)
SKILL_KEYWORDS = {
    "python", "java", "c++", "sql", "html", "css", "javascript", "react", "angular",
    "node", "express", "flask", "django", "pandas", "numpy", "ml", "ai", "nlp",
    "deep learning", "tensorflow", "keras", "pytorch", "data analysis",
    "communication", "leadership", "teamwork", "problem solving",
    "cloud", "aws", "azure", "gcp", "docker", "kubernetes", "linux",
    "jira", "git", "github", "rest api", "fastapi", "power bi", "excel"
}

# ---------- Text Extraction ----------
def extract_text_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
    return text


def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])


def extract_text_from_doc(file_path):
    """Basic support for .doc files using antiword if installed"""
    try:
        import subprocess
        result = subprocess.run(["antiword", file_path], capture_output=True, text=True)
        return result.stdout
    except Exception:
        logger.warning("antiword not available. .doc extraction limited.")
        return ""


def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    return text.strip()

# ---------- Field Extraction ----------
def extract_name(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return None


def extract_email(text):
    match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    return match.group(0) if match else None


def extract_phone(text):
    match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}', text)
    return match.group(0) if match else None


def extract_skills(text):
    text_lower = text.lower()

    # Expanded keyword set (add or pull dynamically from job_metadata.pkl)
    SKILL_KEYWORDS = {
        "python", "java", "c++", "sql", "html", "css", "javascript", "react", "node",
        "flask", "django", "tensorflow", "pytorch", "ml", "ai", "nlp",
        "seo", "sem", "ppc", "digital marketing", "social media", "email marketing",
        "content marketing", "campaign management", "google ads", "meta ads",
        "linkedin ads", "hubspot", "salesforce", "crm", "analytics",
        "google analytics", "lead generation", "branding", "strategy",
        "leadership", "team management", "market research", "consumer insights",
        "copywriting", "ad design", "communication", "presentation", "negotiation",
        "project management", "microsoft office", "excel", "power bi", "jira", "git"
    }

    # Token-based detection
    tokens = [t.text.lower() for t in nlp(text) if not t.is_stop]
    found = set()

    for skill in SKILL_KEYWORDS:
        if skill in text_lower:
            found.add(skill)

    # Multi-word detection (e.g., "digital marketing", "social media marketing")
    multiword_skills = [s for s in SKILL_KEYWORDS if " " in s]
    for skill in multiword_skills:
        if skill in text_lower:
            found.add(skill)

    return sorted(found)



def extract_education(text):
    edu_keywords = [
        "bachelor", "master", "phd", "b.tech", "m.tech", "b.sc", "m.sc",
        "mba", "bca", "mca", "bcom", "mcom", "degree", "diploma"
    ]
    found = [w for w in edu_keywords if w in text.lower()]
    return sorted(set(found))


import re

def extract_experience(text):
    """
    Extract structured work experience data from resume text.
    """
    experience = []
    exp_section = re.search(r'(?i)(professional experience|work experience)(.*?)(education|projects|certifications|achievements|skills|$)', text, re.DOTALL)
    
    if not exp_section:
        return experience
    
    exp_text = exp_section.group(2)

    # Split each role block (e.g., "Marketing Manager | Hindustan Unilever Ltd. ...")
    role_blocks = re.split(r'(?=\b[A-Z][a-zA-Z\s]+ \| )', exp_text)

    for block in role_blocks:
        if len(block.strip()) < 30:
            continue
        
        role_match = re.search(r'^([A-Z][a-zA-Z\s]+)\s*\|\s*([A-Za-z&\s\.]+)', block)
        duration_match = re.search(r'([A-Za-z]{3,9}\s*\d{4})\s*[-–]\s*([A-Za-z]{3,9}\s*\d{4}|Present)', block)
        
        role = role_match.group(1).strip() if role_match else ""
        company = role_match.group(2).strip() if role_match else ""
        duration = f"{duration_match.group(1)} - {duration_match.group(2)}" if duration_match else ""
        
        # Description: everything after the duration
        desc = ""
        if duration_match:
            desc = block[duration_match.end():].strip()
        elif role_match:
            desc = block[role_match.end():].strip()
        
        if role:
            experience.append({
                "role": role,
                "company": company,
                "duration": duration,
                "description": re.sub(r'\s+', ' ', desc)
            })
    
    return experience



# ---------- Main Parser ----------
def parse_resume(resume_path):
    text = ""
    if resume_path.endswith(".pdf"):
        import pdfplumber
        with pdfplumber.open(resume_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    elif resume_path.endswith(".docx"):
        from docx import Document
        doc = Document(resume_path)
        for para in doc.paragraphs:
            text += para.text + "\n"

    skills = extract_skills(text)   # your existing skill extractor
    education = extract_education(text)
    experience = extract_experience(text)

    return {
        "raw_text": text.strip(),
        "skills": skills,
        "education": education,
        "experience": experience
    }



# ---------- Local Test ----------
if __name__ == "__main__":
    sample_file = "uploads/Priya Nair.pdf"
    result = parse_resume(sample_file)
    print("\n✅ Parsed Resume Data:")
    for key, val in result.items():
        print(f"{key}: {val}")
